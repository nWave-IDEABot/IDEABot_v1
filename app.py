import json
import os
import requests
import pandas as pd
import cloudant
from flask_basicauth import BasicAuth
from cloudant import Cloudant
from docx import Document
import time
from flask import Flask , request, make_response , render_template, session, g
from cloudant.error import CloudantException
from cloudant.result import Result, ResultByKey,QueryResult

user= "9ef80d52-30b9-4e45-833a-75059db0825c-bluemix"
password= "79d7c792bb13072da32aecc7dc3777e28780d6d11619795d54d717d2abbd62e5"
host= "9ef80d52-30b9-4e45-833a-75059db0825c-bluemix.cloudant.com"
url = 'https://' + host
client = Cloudant(user, password, url=url, connect=True)    

app = Flask(__name__)

app.config['BASIC_AUTH_USERNAME'] = 'admin'
app.config['BASIC_AUTH_PASSWORD'] = 'admin'
basic_auth = BasicAuth(app)

@app.route('/admin')
@basic_auth.required
def admin():
    #OUTPUTDB
    session = client.session()
    db = client['input-dataset']

    #INPUTDB
    db2= client['nwaveoutput']
    query1 = cloudant.query.Query(db2,selector={"admin-flag":0})
    time.sleep(1)
    zeroflagged = QueryResult(query1)
    print(zeroflagged)
    query2 = cloudant.query.Query(db2,selector={"admin-flag":1})
    time.sleep(1)
    flagged = QueryResult(query2)
    print(flagged)
   
    return render_template('admin.html',approved=flagged,notApproved=zeroflagged)

@app.route('/updateDB/<id>')
def updateDB(id):
    session = client.session()
    db = client['nwaveoutput']
    doc = db[id]
    doc['admin-flag']=1
    doc.save()
    db2=client['input-dataset']
    new_doc=db2.create_document(doc)
    if new_doc.exists():
        print('SUCCESS!!')
    return admin()

@app.route('/rejectData/<id>')
def rejectData(id):
    session = client.session()
    db = client['nwaveoutput']
    doc = db[id]
    doc['admin-flag']=-1
    doc.save()
    return admin()

@app.route('/getop/<sessionId>')
def getop(sessionId):
    session = client.session()
    db = client['nwaveoutput']
    query = cloudant.query.Query(db,selector={"sessionId": sessionId})
    query_result = QueryResult(query)
    #query_result = db.get_query_result(selector)
    time.sleep(1)
    generate_docx(query_result)
    print(query_result)
    return render_template('output.html',weightage=query_result)
    #except:
    #   return "Sorry something went wrong"

@app.route('/feedbackPositive/<id>')
def positiveFeedback(id):
    session = client.session()
    db = client['nwaveoutput']
    doc = db[id]
    doc['feedback']="positive"
    sessionId=doc['sessionId']
    doc.save()
    return getop(sessionId)

@app.route('/negativePositive/<id>')
def negativeFeedback(id):
    session = client.session()
    db = client['nwaveoutput']
    doc = db[id]
    doc['feedback']="negative"
    sessionId=doc['sessionId']
    doc.save()
    return getop(sessionId)

@app.route('/docx')
def download_docx():
    with open("static/estimate.docx", 'rb') as f:
        body = f.read()
    response = make_response(body)
    response.headers["Content-Disposition"] = "attachment; filename=estimate.docx"
    return response    

def generate_docx(query_res):
    document = Document("static/template.docx")      
    for doc in query_res:
        document.add_heading("Interface Details:",level=2)
        document.add_paragraph("Product           : " + doc['disp-product'])
        document.add_paragraph("Source Protocol   : " + doc['disp-srcP'])
        document.add_paragraph("Source Msg Format : " + doc['disp-srcF'])
        document.add_paragraph("Target Protocol   : " + doc['disp-targetP'])
        document.add_paragraph("Target Msg format :" + doc['disp-targetF'])
        
        document.add_heading("Effort Details:",level=2)       
        document.add_paragraph("Design:" + str(doc['weightage']*0.25))
        document.add_paragraph("Test:" + str(doc['weightage']*0.25))
        document.add_paragraph("Development:" + str(doc['weightage']*0.5))
        document.add_paragraph("Total Effort:" + str(doc['weightage']))
        
    document.add_heading("Consolidated Details:",level=2)    
    table = document.add_table(rows=1, cols=6, style='TableGrid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Interface Details'
    hdr_cells[1].text = 'Effort'
    time.sleep(1)
    for doc in query_res:
        row_cells = table.add_row().cells        
        row_cells[0].text = doc['disp-product']
        row_cells[1].text = str(doc['weightage'])
    document.add_paragraph("")    
    document.add_paragraph("© 2018 Cognizant EAS-IPM.")
    document.save("static/estimate.docx")
    return document

port = os.getenv('VCAP_APP_PORT', '5000')
if __name__ == "__main__":
       	app.run(host='0.0.0.0', port=int(port), use_reloader=True, debug=True)
